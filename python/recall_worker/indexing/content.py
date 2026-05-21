from __future__ import annotations

import tempfile
from dataclasses import dataclass
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont


@dataclass(slots=True)
class TextChunk:
    text: str
    chunk_index: int
    chunk_type: str
    page_number: int | None = None
    start_ms: int | None = None
    end_ms: int | None = None


def chunk_text(
    text: str,
    *,
    chunk_type: str,
    words_per_chunk: int = 800,
    overlap_words: int = 120,
    page_number: int | None = None,
) -> list[TextChunk]:
    words = text.split()
    if not words:
        return []

    chunks: list[TextChunk] = []
    start = 0
    chunk_index = 0
    step = max(1, words_per_chunk - overlap_words)
    while start < len(words):
        slice_words = words[start : start + words_per_chunk]
        if not slice_words:
            break
        chunks.append(
            TextChunk(
                text=" ".join(slice_words).strip(),
                chunk_index=chunk_index,
                chunk_type=chunk_type,
                page_number=page_number,
            )
        )
        chunk_index += 1
        start += step
    return chunks


def extract_document_chunks(document_path: Path, ocr_engine) -> tuple[list[TextChunk], str | None]:
    suffix = document_path.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf_chunks(document_path, ocr_engine)
    if suffix == ".docx":
        return _extract_docx_chunks(document_path), None
    if suffix == ".txt":
        return _extract_txt_chunks(document_path), None
    raise ValueError(f"Unsupported document type: {document_path.suffix}")


def extract_audio_chunks(audio_path: Path, transcription_engine) -> tuple[list[TextChunk], int | None]:
    segments = transcription_engine.transcribe(audio_path)
    if not segments:
        return [], None

    chunks: list[TextChunk] = []
    current_text: list[str] = []
    current_start: int | None = None
    current_end: int | None = None
    chunk_index = 0

    for segment in segments:
        start_ms = int(segment["startMs"])
        end_ms = int(segment["endMs"])
        text = str(segment["text"]).strip()
        if not text:
            continue

        if current_start is None:
            current_start = start_ms
        if current_end is None:
            current_end = end_ms

        proposed_end = end_ms
        proposed_duration = proposed_end - current_start
        if current_text and proposed_duration > 45000:
            chunks.append(
                TextChunk(
                    text=" ".join(current_text).strip(),
                    chunk_index=chunk_index,
                    chunk_type="transcript",
                    start_ms=current_start,
                    end_ms=current_end,
                )
            )
            chunk_index += 1
            current_text = [text]
            current_start = start_ms
            current_end = end_ms
            continue

        current_text.append(text)
        current_end = end_ms

    if current_text and current_start is not None and current_end is not None:
        chunks.append(
            TextChunk(
                text=" ".join(current_text).strip(),
                chunk_index=chunk_index,
                chunk_type="transcript",
                start_ms=current_start,
                end_ms=current_end,
            )
        )

    duration_ms = max((int(segment["endMs"]) for segment in segments), default=0)
    return chunks, duration_ms or None


def render_document_preview(
    document_path: Path,
    preview_path: Path,
    *,
    max_size: int,
    chunks: list[TextChunk],
) -> Path:
    if preview_path.exists():
        return preview_path

    preview_path.parent.mkdir(parents=True, exist_ok=True)
    suffix = document_path.suffix.lower()
    if suffix == ".pdf":
        image = _render_pdf_preview_image(document_path)
    else:
        image = _render_text_document_preview_image(document_path, chunks)

    try:
        image.thumbnail((max_size, max_size), Image.Resampling.LANCZOS)
        image.save(preview_path, format="JPEG", quality=90)
    finally:
        image.close()
    return preview_path


def _extract_txt_chunks(document_path: Path) -> list[TextChunk]:
    raw_bytes = document_path.read_bytes()
    for encoding in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            text = raw_bytes.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    else:
        text = raw_bytes.decode("utf-8", errors="replace")
    return chunk_text(text, chunk_type="document")


def _extract_docx_chunks(document_path: Path) -> list[TextChunk]:
    from docx import Document  # type: ignore

    document = Document(str(document_path))
    text = "\n".join(paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip())
    return chunk_text(text, chunk_type="document")


def _extract_pdf_chunks(document_path: Path, ocr_engine) -> tuple[list[TextChunk], str | None]:
    from pypdf import PdfReader  # type: ignore

    reader = PdfReader(str(document_path))
    page_texts: list[tuple[int, str]] = []
    for page_number, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            page_texts.append((page_number, text))

    total_text = " ".join(text for _, text in page_texts).strip()
    if len(total_text) >= 80:
        chunks: list[TextChunk] = []
        for page_number, text in page_texts:
            chunks.extend(chunk_text(text, chunk_type="document", page_number=page_number))
        return chunks, None

    return _extract_pdf_chunks_with_ocr(document_path, ocr_engine), "ocr_fallback"


def _extract_pdf_chunks_with_ocr(document_path: Path, ocr_engine) -> list[TextChunk]:
    import pypdfium2 as pdfium  # type: ignore

    pdf = pdfium.PdfDocument(str(document_path))
    chunks: list[TextChunk] = []
    try:
        ocr_engine.ensure_ready()
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_root = Path(temp_dir)
            for page_index in range(len(pdf)):
                page = pdf[page_index]
                bitmap = page.render(scale=2.0)
                pil_image = bitmap.to_pil()
                image_path = temp_root / f"page-{page_index + 1}.png"
                pil_image.save(image_path)
                text = ocr_engine.extract_text(image_path)
                if text.strip():
                    chunks.extend(
                        chunk_text(
                            text,
                            chunk_type="document",
                            page_number=page_index + 1,
                        )
                    )
    finally:
        pdf.close()
    return chunks


def _render_pdf_preview_image(document_path: Path) -> Image.Image:
    import pypdfium2 as pdfium  # type: ignore

    pdf = pdfium.PdfDocument(str(document_path))
    try:
        page = pdf[0]
        bitmap = page.render(scale=1.4)
        return bitmap.to_pil().convert("RGB")
    finally:
        pdf.close()


def _render_text_document_preview_image(
    document_path: Path,
    chunks: list[TextChunk],
) -> Image.Image:
    canvas = Image.new("RGB", (900, 1160), color=(241, 247, 250))
    draw = ImageDraw.Draw(canvas)

    draw.rounded_rectangle((48, 36, 852, 1112), radius=34, fill=(252, 254, 255))
    draw.polygon([(680, 36), (852, 36), (852, 208)], fill=(228, 238, 244))
    draw.line((680, 36, 852, 208), fill=(214, 227, 236), width=2)

    title_font = _load_font(36, bold=True)
    meta_font = _load_font(22, bold=True)
    body_font = _load_font(22)

    extension = document_path.suffix.replace(".", "").upper() or "DOC"
    draw.rounded_rectangle((96, 92, 208, 146), radius=26, fill=(94, 201, 227))
    draw.text((124, 104), extension, fill=(7, 33, 46), font=meta_font)

    filename = document_path.stem[:44]
    draw.text((96, 184), filename, fill=(24, 54, 70), font=title_font)

    sample_text = " ".join(chunk.text.strip() for chunk in chunks[:2]).strip()
    if not sample_text:
        sample_text = document_path.stem.replace("_", " ").replace("-", " ").strip() or "Document preview"

    lines = _wrap_text(draw, sample_text, body_font, max_width=708)
    y = 272
    line_height = 34
    for line in lines[:20]:
        draw.text((104, y), line, fill=(67, 93, 108), font=body_font)
        y += line_height
        if y > 1016:
            break

    return canvas


def _load_font(size: int, *, bold: bool = False) -> ImageFont.ImageFont:
    candidates = [
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size=size)
        except OSError:
            continue
    return ImageFont.load_default()


def _wrap_text(
    draw: ImageDraw.ImageDraw,
    text: str,
    font: ImageFont.ImageFont,
    *,
    max_width: int,
) -> list[str]:
    words = text.split()
    if not words:
        return []

    lines: list[str] = []
    current = words[0]
    for word in words[1:]:
        candidate = f"{current} {word}"
        if draw.textlength(candidate, font=font) <= max_width:
            current = candidate
        else:
            lines.append(current)
            current = word
    lines.append(current)
    return lines
